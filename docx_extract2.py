from docx import Document

docx_path = r"C:\Users\xa751\Downloads\Chrome 瀏覽器隱藏秘技與效能工具：全面簡報指令.docx"
doc = Document(docx_path)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text:
        print(f"[{i}][{para.style.name}] {text}")

print("\n========= TABLES =========")
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti+1} ---")
    for ri, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print(f"  Row{ri}: {cells}")
